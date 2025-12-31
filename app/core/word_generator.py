# ==========================================
# GENERADOR DE DOCUMENTOS WORD
# ==========================================

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os


def aplicar_estilos_word(paragraph, run, config_estilo):
    """Aplica formato granular al objeto Word"""
    # 1. Fuente y Tamaño
    run.font.name = config_estilo.get('fuente', 'Arial')
    run.font.size = Pt(config_estilo.get('tamano', 11))
    
    # 2. Estilos específicos por Acción
    accion = config_estilo.get('accion', 'Texto Normal')
    if accion == "Título":
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif accion == "Dato Clave":
        run.bold = True
        run.font.color.rgb = RGBColor(0, 50, 150)
    elif accion == "Subtítulo":
        run.font.color.rgb = RGBColor(80, 80, 80)
    
    # 3. Alineación
    align_map = {
        "Izquierda": WD_ALIGN_PARAGRAPH.LEFT,
        "Centrado": WD_ALIGN_PARAGRAPH.CENTER,
        "Derecha": WD_ALIGN_PARAGRAPH.RIGHT,
        "Justificado": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    paragraph.alignment = align_map.get(config_estilo.get('alineacion', 'Izquierda'))
    
    # 4. Interlineado
    spacing_map = {"Simple": 1.0, "1.15": 1.15, "1.5": 1.5, "Doble": 2.0}
    val_spacing = spacing_map.get(config_estilo.get('interlineado', 'Simple'))
    paragraph.paragraph_format.line_spacing = val_spacing


def agregar_heading_toc(doc, titulo_toc, nivel_word, config_toc):
    """Agrega un heading de TOC al documento Word"""
    h = doc.add_heading(titulo_toc, level=nivel_word)
    for run in h.runs:
        run.font.name = config_toc.get('fuente', 'Arial')
        run.font.size = Pt(config_toc.get('tamano', 14) if nivel_word == 1 else config_toc.get('tamano', 14)-2)
        run.font.color.rgb = RGBColor(60, 60, 60)
    return h


def agregar_separador_pagina(doc, num_pag):
    """Agrega un separador de página al documento Word"""
    txt_sep = f"--- Página {num_pag} ---"
    p = doc.add_paragraph(txt_sep)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    return txt_sep


def agregar_texto_resaltado(doc, texto_final, conf):
    """Agrega texto resaltado al documento Word con el estilo apropiado"""
    style = None
    if conf.get('lista') == "Bullets (•)": 
        style = 'List Bullet'
    elif conf.get('lista') == "Numerada (1.)": 
        style = 'List Number'
    
    if conf['accion'] == "Título":
        p = doc.add_heading(texto_final, level=2)
    elif conf['accion'] == "Subtítulo":
        p = doc.add_heading(texto_final, level=3)
    else:
        p = doc.add_paragraph(texto_final, style=style)
    
    if p.runs: 
        run = p.runs[0]
    else: 
        run = p.add_run(texto_final)
        
    aplicar_estilos_word(p, run, conf)
    return p


def guardar_documento_word(doc):
    """Guarda el documento Word en carpeta export y retorna buffer"""
    # Guardado automático en carpeta export
    if not os.path.exists("export"):
        os.makedirs("export")
    doc.save(os.path.join("export", "resumen_generado.docx"))

    # Retornar buffer para descarga
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
